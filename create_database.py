# create_database.py
import os
from pathlib import Path
from typing import List
import pinecone
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Pinecone

DATA_PATH = "data/"
INDEX_NAME = os.environ.get("PINECONE_INDEX", "portfolio-rag")
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
DIMENSION = 384

def list_documents() -> List[str]:
    Path(DATA_PATH).mkdir(parents=True, exist_ok=True)
    files = [
        f for f in os.listdir(DATA_PATH)
        if f.lower().endswith((".txt", ".md", ".docx"))
    ]
    return files

def _load_file_as_documents(path: str, filename: str) -> List[Document]:
    ext = Path(path).suffix.lower()
    text = ""
    if ext in [".txt", ".md"]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
        except Exception:
            text = ""
    if not text or not text.strip():
        return []
    # return a single Document with metadata source
    return [Document(page_content=text, metadata={"source": filename})]

def generate_data_store() -> bool:
    """
    Rebuilds Pinecone index from all files in DATA_PATH.
    Returns True on success, False otherwise.
    """
    # Ensure environment variables set
    api_key = os.environ.get("PINECONE_API_KEY")
    env = os.environ.get("PINECONE_ENV")
    if not api_key or not env:
        print("Missing PINECONE_API_KEY or PINECONE_ENV environment variable.")
        return False

    # init pinecone
    pinecone.init(api_key=api_key, environment=env)

    # delete existing index (so deleted files are forgotten)
    existing = pinecone.list_indexes()
    if INDEX_NAME in existing:
        try:
            pinecone.delete_index(INDEX_NAME)
        except Exception as e:
            print("Warning: couldn't delete index:", e)

    # create index
    try:
        pinecone.create_index(INDEX_NAME, dimension=DIMENSION, metric="cosine")
    except Exception as e:
        # index creation sometimes raises if just deleted; try again later or continue
        print("Index create warning:", e)

    # collect documents
    docs = []
    for fname in list_documents():
        path = os.path.join(DATA_PATH, fname)
        docs.extend(_load_file_as_documents(path, fname))

    if not docs:
        print("No valid documents found to index.")
        return True  # Not an error — just nothing to index

    # chunk documents
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    splitted = text_splitter.split_documents(docs)

    # embeddings
    embedding = HuggingFaceEmbeddings(model_name=EMBED_MODEL)

    # create/ upsert to Pinecone via LangChain helper
    try:
        Pinecone.from_documents(splitted, embedding, index_name=INDEX_NAME)
    except Exception as e:
        print("Failed to push documents to Pinecone:", e)
        return False

    print(f"Indexed {len(splitted)} chunks into Pinecone index '{INDEX_NAME}'.")
    return True
